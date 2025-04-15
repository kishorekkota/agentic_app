from docx import Document
from io import BytesIO

def generate_doc(chat_history):
    doc = Document()
    doc.add_heading('Job Description', 0)

    doc.add_paragraph(f"{chat_history.get('message')}")
    doc.add_paragraph("")  # Add a line break between messages

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer